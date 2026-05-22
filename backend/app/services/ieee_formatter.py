import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, ns

def set_margins(doc):
    """Sets IEEE standard margins."""
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(0.63)
        section.right_margin = Inches(0.63)

def add_page_numbering(doc):
    """Adds standard page numbering to the footer using XML elements."""
    for section in doc.sections:
        footer = section.footer
        paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        run = paragraph.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(ns.qn('w:fldCharType'), 'begin')
        
        instrText = OxmlElement('w:instrText')
        instrText.set(ns.qn('xml:space'), 'preserve')
        instrText.text = "PAGE"
        
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(ns.qn('w:fldCharType'), 'end')
        
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)

def set_styles(doc):
    """Configures global IEEE styles (Times New Roman, sizes, headings)."""
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(10)
    
    if 'Title' in doc.styles:
        title_style = doc.styles['Title']
        title_style.font.name = 'Times New Roman'
        title_style.font.size = Pt(24)
        title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
    if 'Heading 1' in doc.styles:
        h1_style = doc.styles['Heading 1']
        h1_style.font.name = 'Times New Roman'
        h1_style.font.size = Pt(10)
        h1_style.font.small_caps = True
        h1_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        h1_style.paragraph_format.space_before = Pt(12)
        h1_style.paragraph_format.space_after = Pt(6)

def generate_ieee_docx(parsed_data: dict, output_path: str):
    """Generates a formatted DOCX document strictly adhering to IEEE guidelines."""
    doc = docx.Document()
    
    set_margins(doc)
    set_styles(doc)
    add_page_numbering(doc)
    
    # Title
    if parsed_data.get('title'):
        doc.add_paragraph(parsed_data['title'], style='Title')
        
    # Authors
    if parsed_data.get('authors'):
        author_para = doc.add_paragraph()
        author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        author_para.add_run(", ".join(parsed_data['authors'])).font.size = Pt(11)
        
    # Abstract
    if parsed_data.get('abstract'):
        doc.add_paragraph() # Spacer
        abs_para = doc.add_paragraph()
        r1 = abs_para.add_run("Abstract—")
        r1.bold = True
        r1.italic = True
        r2 = abs_para.add_run(parsed_data['abstract'])
        r2.bold = True
        
    # Sections (with Roman Numeral headings)
    roman_numerals = ["I", "II", "III", "IV", "V", "VI", "VII", "VIII", "IX", "X"]
    for idx, section in enumerate(parsed_data.get('sections', [])):
        heading = section.get('heading', '')
        if heading:
            roman = roman_numerals[idx] if idx < len(roman_numerals) else str(idx + 1)
            doc.add_paragraph(f"{roman}. {heading}", style='Heading 1')
            
        content = section.get('content', '')
        if content:
            p = doc.add_paragraph(content)
            p.paragraph_format.first_line_indent = Inches(0.14)
            p.paragraph_format.space_after = Pt(0)
            
    # References
    if parsed_data.get('references'):
        doc.add_paragraph("REFERENCES", style='Heading 1')
        for idx, ref in enumerate(parsed_data['references']):
            p = doc.add_paragraph(f"[{idx+1}] {ref}")
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.first_line_indent = Inches(-0.25)

    doc.save(output_path)
    return output_path
