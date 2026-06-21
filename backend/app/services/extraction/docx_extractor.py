"""DOCX Extractor — V2 with ZIP architecture.

Extracts text, headings, styles, tables, images, captions, and references
by directly parsing the DOCX ZIP structure (word/document.xml, word/styles.xml,
word/media/*, relationship files).
"""

import os
import re
import time
import zipfile
import xml.etree.ElementTree as ET
from typing import Dict, List, Optional, Tuple

from app.services.extraction.base import BaseExtractor
from app.services.extraction.report import (
    ExtractionResult,
    ExtractionMetadata,
    ExtractedSection,
    ExtractedFigure,
    ExtractedTable,
    ExtractedReference,
)

# Word Open XML namespace
NS = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
    "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
    "pic": "http://schemas.openxmlformats.org/drawingml/2006/picture",
    "rel": "http://schemas.openxmlformats.org/package/2006/relationships",
}

HEADING_LABELS = {
    "abstract": "abstract",
    "introduction": "introduction",
    "related work": "related_work",
    "literature review": "related_work",
    "methodology": "methodology",
    "methods": "methods",
    "method": "methods",
    "proposed method": "methods",
    "proposed approach": "methods",
    "experiments": "experiments",
    "experimental setup": "experiments",
    "evaluation": "experiments",
    "results": "results",
    "results and discussion": "discussion",
    "discussion": "discussion",
    "conclusion": "conclusion",
    "conclusions": "conclusion",
    "acknowledgments": "acknowledgments",
    "acknowledgements": "acknowledgments",
    "references": "references",
    "bibliography": "references",
    "appendix": "appendix",
}


def _classify_heading(text: str) -> str:
    """Classify a heading into a section type."""
    norm = text.lower().strip()
    for key, label in HEADING_LABELS.items():
        if key in norm or norm in key:
            return label
    return "other"


def _is_heading(paragraph) -> bool:
    """Check if a paragraph is a heading based on style."""
    if not paragraph.style:
        return False
    style_name = paragraph.style.name.lower()
    return "heading" in style_name or style_name.startswith("title")


def _get_heading_level(paragraph) -> int:
    """Extract heading level from style name."""
    if not paragraph.style:
        return 1
    match = re.search(r"heading\s*(\d)", paragraph.style.name, re.IGNORECASE)
    if match:
        return int(match.group(1))
    return 1


class DOCXExtractor(BaseExtractor):
    """Extract structured content from DOCX files using ZIP architecture."""

    @property
    def supported_extensions(self) -> list[str]:
        return [".docx"]

    def extract(self, file_path: str) -> ExtractionResult:
        start = time.time()
        result = ExtractionResult()
        result.metadata = ExtractionMetadata(
            file_type="docx",
            file_size_bytes=os.path.getsize(file_path),
            parser_used="docx_extractor_v2",
        )

        try:
            # Phase 1: Parse main document content
            self._extract_document_content(file_path, result)

            # Phase 2: Extract styles information
            self._extract_styles(file_path, result)

            # Phase 3: Extract media references
            self._extract_media(file_path, result)

            # Phase 4: Build raw text
            result.raw_text = self._build_raw_text(result)

        except Exception as e:
            result.metadata.warnings.append(f"Extraction error: {str(e)}")

        result.metadata.processing_time_ms = (time.time() - start) * 1000
        return result

    def _extract_document_content(self, file_path: str, result: ExtractionResult):
        """Parse word/document.xml for text, headings, tables, and figures."""
        try:
            import docx
            doc = docx.Document(file_path)
        except Exception:
            return

        title = ""
        authors = []
        abstract = ""
        sections: List[ExtractedSection] = []
        references: List[ExtractedReference] = []
        figures: List[ExtractedFigure] = []
        tables: List[ExtractedTable] = []

        current_section_type = "general"
        current_heading = ""
        section_order = 0
        ref_index = 0
        abstract_pattern = re.compile(r"^(abstract|summary)$", re.IGNORECASE)
        reference_pattern = re.compile(r"^(references|bibliography|works cited)$", re.IGNORECASE)
        figure_pattern = re.compile(r"^(fig\.|figure)\s*\d+[:\.\-]", re.IGNORECASE)
        ref_bracket_pattern = re.compile(r"\[(\d+(?:,\s*\d+)*)\]")

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            style_name = para.style.name.lower() if para.style else ""

            # Title detection
            if not title and ("title" in style_name or len(text) < 150):
                title = text
                continue

            # Heading detection
            if _is_heading(para) or text.isupper():
                if abstract_pattern.match(text):
                    current_section_type = "abstract"
                    continue
                elif reference_pattern.match(text):
                    current_section_type = "references"
                    continue
                else:
                    current_section_type = "body"
                    current_heading = text
                    section_order += 1
                    sections.append(ExtractedSection(
                        heading=current_heading,
                        content="",
                        level=_get_heading_level(para),
                        order=section_order,
                    ))
                    continue

            # Figure caption detection
            if figure_pattern.match(text) or "caption" in style_name:
                figures.append(ExtractedFigure(caption=text))
                continue

            # Content routing
            if current_section_type == "abstract":
                abstract += text + "\n"
            elif current_section_type == "references":
                # Try to extract individual references
                ref_entries = self._split_references(text)
                for ref_text in ref_entries:
                    ref_index += 1
                    doi = self._extract_doi(ref_text)
                    references.append(ExtractedReference(
                        raw_text=ref_text,
                        index=ref_index,
                        doi=doi,
                    ))
            elif current_section_type == "body":
                if sections:
                    sections[-1].content += text + "\n"
                else:
                    section_order += 1
                    sections.append(ExtractedSection(
                        heading="",
                        content=text + "\n",
                        order=section_order,
                    ))
            else:
                if sections:
                    sections[-1].content += text + "\n"
                else:
                    section_order += 1
                    sections.append(ExtractedSection(
                        heading="",
                        content=text + "\n",
                        order=section_order,
                    ))
                    current_section_type = "body"

        # Extract tables
        for i, table in enumerate(doc.tables):
            table_rows = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_rows.append(row_data)
            if table_rows:
                caption = f"Table {i + 1}"
                tables.append(ExtractedTable(
                    caption=caption,
                    rows=table_rows,
                    order=i + 1,
                ))

        result.title = title
        result.authors = authors
        result.abstract = abstract.strip()
        result.sections = sections
        result.references = references
        result.figures = figures
        result.tables = tables
        result.metadata.has_tables = len(tables) > 0
        result.metadata.has_images = len(figures) > 0

    def _extract_styles(self, file_path: str, result: ExtractionResult):
        """Extract style information from word/styles.xml."""
        try:
            with zipfile.ZipFile(file_path, "r") as z:
                if "word/styles.xml" in z.namelist():
                    styles_xml = z.read("word/styles.xml")
                    # Parse and count styles
                    root = ET.fromstring(styles_xml)
                    style_count = len(root.findall(".//w:style", NS))
                    if style_count > 0:
                        result.metadata.styles_extracted = True
        except Exception:
            pass

    def _extract_media(self, file_path: str, result: ExtractionResult):
        """Extract media file references from the DOCX ZIP."""
        try:
            with zipfile.ZipFile(file_path, "r") as z:
                media_files = [f for f in z.namelist() if f.startswith("word/media/")]
                if media_files:
                    result.metadata.has_images = True
                    for mf in media_files:
                        ext = os.path.splitext(mf)[1].lower()
                        mime_map = {
                            ".png": "image/png",
                            ".jpg": "image/jpeg",
                            ".jpeg": "image/jpeg",
                            ".gif": "image/gif",
                            ".emf": "image/x-emf",
                            ".wmf": "image/x-wmf",
                        }
                        mime = mime_map.get(ext, "image/unknown")
                        caption = os.path.basename(mf)
                        result.figures.append(ExtractedFigure(
                            caption=caption,
                            path=mf,
                            mime_type=mime,
                        ))
        except Exception:
            pass

    def _split_references(self, text: str) -> List[str]:
        """Split a reference block into individual references."""
        # Try numbered references: [1] Author, Title...
        numbered = re.split(r"\[\d+\]\s*", text)
        if len(numbered) > 1:
            return [r.strip() for r in numbered if r.strip()]

        # Try line-by-line
        lines = text.split("\n")
        refs = []
        current = ""
        for line in lines:
            line = line.strip()
            if not line:
                if current:
                    refs.append(current)
                    current = ""
                continue
            if re.match(r"^\d+[\.\)]\s", line):
                if current:
                    refs.append(current)
                current = re.sub(r"^\d+[\.\)]\s*", "", line)
            else:
                current += " " + line
        if current:
            refs.append(current)
        return refs if refs else [text]

    def _extract_doi(self, text: str) -> Optional[str]:
        """Extract DOI from reference text."""
        doi_match = re.search(r"10\.\d{4,}/[^\s\)\],;]+", text)
        if doi_match:
            return doi_match.group(0).rstrip(".,;)")
        return None

    def _build_raw_text(self, result: ExtractionResult) -> str:
        """Build a single raw text string from all extracted content."""
        parts = []
        if result.title:
            parts.append(f"Title: {result.title}")
        if result.authors:
            parts.append(f"Authors: {', '.join(result.authors)}")
        if result.abstract:
            parts.append(f"Abstract: {result.abstract}")
        for sec in result.sections:
            if sec.heading:
                parts.append(f"\n## {sec.heading}\n{sec.content}")
            else:
                parts.append(sec.content)
        if result.references:
            parts.append("\nReferences:")
            for ref in result.references:
                parts.append(f"[{ref.index}] {ref.raw_text}")
        return "\n".join(parts)
