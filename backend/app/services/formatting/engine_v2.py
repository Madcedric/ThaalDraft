"""Formatting Engine — v2.

Works exclusively from StructuredManuscript data.
Never operates on raw text. Each template produces a properly formatted DOCX.
"""
import os
import time
import logging
from typing import Any, Dict, List, Optional

from app.services.manuscript.model import (
    StructuredManuscript,
    ManuscriptSection,
    SectionType,
    Reference,
    Table,
    Figure,
)

logger = logging.getLogger(__name__)


class FormatConfig:
    """Configuration for a journal format template."""
    def __init__(
        self,
        name: str,
        body_font: str = "Times New Roman",
        body_size: int = 10,
        title_size: int = 24,
        heading_sizes: Optional[List[int]] = None,
        margins: Optional[Dict[str, float]] = None,
        line_spacing: float = 1.0,
        two_column: bool = False,
        references_label: str = "References",
        citation_format: str = "numbered",
        keywords_label: str = "Keywords",
        requires_keywords: bool = False,
        abstract_label: str = "Abstract",
        hanging_indent: bool = False,
    ):
        self.name = name
        self.body_font = body_font
        self.body_size = body_size
        self.title_size = title_size
        self.heading_sizes = heading_sizes or [12, 11]
        self.margins = margins or {"top": 1.0, "bottom": 1.0, "left": 1.0, "right": 1.0}
        self.line_spacing = line_spacing
        self.two_column = two_column
        self.references_label = references_label
        self.citation_format = citation_format
        self.keywords_label = keywords_label
        self.requires_keywords = requires_keywords
        self.abstract_label = abstract_label
        self.hanging_indent = hanging_indent


TEMPLATE_CONFIGS = {
    "ieee": FormatConfig(
        name="IEEE",
        body_size=10,
        title_size=24,
        heading_sizes=[10, 10],
        margins={"top": 0.75, "bottom": 1.0, "left": 0.63, "right": 0.63},
        two_column=True,
        citation_format="numbered",
    ),
    "acm": FormatConfig(
        name="ACM",
        body_size=10,
        title_size=14,
        heading_sizes=[12, 11],
        margins={"top": 0.75, "bottom": 1.0, "left": 0.75, "right": 0.75},
        two_column=True,
        requires_keywords=True,
        keywords_label="Keywords",
    ),
    "springer": FormatConfig(
        name="Springer LNCS",
        body_size=10,
        title_size=14,
        heading_sizes=[12, 11],
        margins={"top": 0.75, "bottom": 1.0, "left": 0.75, "right": 0.75},
        two_column=True,
    ),
    "apa": FormatConfig(
        name="APA 7th Edition",
        body_size=12,
        title_size=14,
        heading_sizes=[14, 12],
        margins={"top": 1.0, "bottom": 1.0, "left": 1.0, "right": 1.0},
        line_spacing=2.0,
        hanging_indent=True,
    ),
    "mla": FormatConfig(
        name="MLA 9th Edition",
        body_size=12,
        title_size=14,
        heading_sizes=[14, 12],
        margins={"top": 1.0, "bottom": 1.0, "left": 1.0, "right": 1.0},
        line_spacing=2.0,
        references_label="Works Cited",
        hanging_indent=True,
    ),
    "nature": FormatConfig(
        name="Nature",
        body_size=10,
        title_size=16,
        heading_sizes=[12, 11],
        margins={"top": 1.0, "bottom": 1.0, "left": 1.0, "right": 1.0},
        two_column=False,
    ),
    "elsevier": FormatConfig(
        name="Elsevier",
        body_size=10,
        title_size=14,
        heading_sizes=[12, 11],
        margins={"top": 1.0, "bottom": 1.0, "left": 1.0, "right": 1.0},
        two_column=False,
        requires_keywords=True,
        keywords_label="Keywords",
        line_spacing=1.5,
    ),
}


def _roman_numeral(idx: int) -> str:
    """Convert 0-based index to Roman numeral."""
    vals = [10, 9, 5, 4, 1]
    syms = ["X", "IX", "V", "IV", "I"]
    result = ""
    for i, v in enumerate(vals):
        while idx >= v:
            result += syms[i]
            idx -= v
    return result or "I"


def _format_reference(ref: Reference, config: FormatConfig) -> str:
    """Format a single reference according to template style."""
    parts = []
    if ref.authors:
        parts.append(", ".join(ref.authors))
    if ref.title:
        parts.append(f'"{ref.title}."')
    if ref.journal:
        parts.append(ref.journal)
    if ref.volume:
        parts.append(f"vol. {ref.volume}")
    if ref.pages:
        parts.append(f"pp. {ref.pages}")
    if ref.year:
        parts.append(str(ref.year))
    if ref.doi:
        parts.append(f"doi: {ref.doi}")

    if not parts:
        return ref.raw_text

    if config.citation_format == "numbered":
        return f"[{ref.index}] {', '.join(parts)}."
    else:
        return f"{'. '.join(parts)}."


def _build_docx(manuscript: StructuredManuscript, template_id: str) -> Any:
    """Build a python-docx Document from a StructuredManuscript."""
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    config = TEMPLATE_CONFIGS.get(template_id, TEMPLATE_CONFIGS["ieee"])

    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(config.margins["top"])
        section.bottom_margin = Inches(config.margins["bottom"])
        section.left_margin = Inches(config.margins["left"])
        section.right_margin = Inches(config.margins["right"])

    style = doc.styles["Normal"]
    style.font.name = config.body_font
    style.font.size = Pt(config.body_size)
    style.paragraph_format.line_spacing = config.line_spacing

    # Title
    if manuscript.title:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(manuscript.title)
        run.bold = True
        run.font.size = Pt(config.title_size)
        run.font.name = config.body_font

    # Authors
    if manuscript.authors:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        author_names = [a.name for a in manuscript.authors]
        run = p.add_run(", ".join(author_names))
        run.font.size = Pt(config.body_size + 1)

        for author in manuscript.authors:
            if author.affiliation:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                aff_run = p.add_run(f"{author.name}, {author.affiliation}")
                aff_run.font.size = Pt(config.body_size - 1)
                aff_run.italic = True
                break

    # Abstract
    if manuscript.abstract:
        doc.add_paragraph()
        p = doc.add_paragraph()
        abs_label = p.add_run(f"{config.abstract_label}: ")
        abs_label.bold = True
        abs_label.italic = True
        abs_text = p.add_run(manuscript.abstract)
        abs_text.italic = True

    # Keywords
    if manuscript.keywords and config.requires_keywords:
        p = doc.add_paragraph()
        kw_label = p.add_run(f"{config.keywords_label}: ")
        kw_label.bold = True
        p.add_run(", ".join(manuscript.keywords))

    # Sections
    roman_numerals = ["I", "II", "III", "IV", "V", "VI", "VII", "VIII", "IX", "X"]
    section_counter = 0

    for sec in manuscript.sections:
        if sec.label in (SectionType.TITLE, SectionType.ABSTRACT, SectionType.KEYWORDS, SectionType.REFERENCES):
            continue

        # Section heading
        if sec.heading:
            heading_text = sec.heading
            if template_id == "ieee" and section_counter < len(roman_numerals):
                heading_text = f"{roman_numerals[section_counter]}. {heading_text}"

            p = doc.add_paragraph()
            heading_size = config.heading_sizes[0] if config.heading_sizes else 12

            if sec.level == 1 and template_id == "ieee":
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            run = p.add_run(heading_text)
            run.bold = True
            run.font.size = Pt(heading_size)
            run.font.name = config.body_font

            if template_id == "ieee" and sec.level == 1:
                run.font.small_caps = True

            section_counter += 1

        # Section content
        if sec.content:
            paragraphs = [p.strip() for p in sec.content.split("\n") if p.strip()]
            for para_text in paragraphs:
                p = doc.add_paragraph(para_text)
                if config.hanging_indent:
                    p.paragraph_format.first_line_indent = Inches(0.5)
                elif template_id == "ieee":
                    p.paragraph_format.first_line_indent = Inches(0.14)

        # Tables in section
        for table in sec.tables:
            if table.headers or table.rows:
                _add_table_to_doc(doc, table)

        # Figures in section
        for figure in sec.figures:
            if figure.caption:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(f"[{figure.caption}]")
                run.italic = True

    # References
    if manuscript.references:
        doc.add_paragraph()
        p = doc.add_paragraph()
        ref_label = p.add_run(config.references_label.upper() if template_id == "ieee" else config.references_label)
        ref_label.bold = True
        if template_id in ("apa", "mla"):
            ref_label.font.size = Pt(14)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        for ref in manuscript.references:
            ref_text = _format_reference(ref, config)
            p = doc.add_paragraph(ref_text)
            if config.hanging_indent:
                p.paragraph_format.left_indent = Inches(0.5)
                p.paragraph_format.first_line_indent = Inches(-0.5)
            elif template_id in ("ieee", "acm", "springer"):
                p.paragraph_format.left_indent = Inches(0.25)
                p.paragraph_format.first_line_indent = Inches(-0.25)

    return doc


def _add_table_to_doc(doc: Any, table: Table) -> None:
    """Add a table to a DOCX document."""
    from docx.shared import Pt

    if not table.headers and not table.rows:
        return

    num_cols = len(table.headers) if table.headers else (len(table.rows[0]) if table.rows else 0)
    if num_cols == 0:
        return

    rows_count = 1 + len(table.rows)
    tbl = doc.add_table(rows=rows_count, cols=num_cols)
    tbl.style = "Table Grid"

    if table.headers:
        for i, header in enumerate(table.headers):
            cell = tbl.rows[0].cells[i]
            cell.text = header
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(9)

    for row_idx, row_data in enumerate(table.rows):
        for col_idx, cell_text in enumerate(row_data):
            if col_idx < num_cols:
                cell = tbl.rows[row_idx + 1].cells[col_idx]
                cell.text = cell_text
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)

    if table.caption:
        p = doc.add_paragraph()
        run = p.add_run(f"Table: {table.caption}")
        run.italic = True
        run.font.size = Pt(9)


def format_manuscript(
    manuscript: StructuredManuscript,
    template_id: str,
    output_dir: str = "exports",
) -> str:
    """Format a StructuredManuscript into a DOCX file. Returns file path."""
    start = time.time()

    if template_id not in TEMPLATE_CONFIGS:
        raise ValueError(f"Unknown template: {template_id}. Available: {list(TEMPLATE_CONFIGS.keys())}")

    doc = _build_docx(manuscript, template_id)

    os.makedirs(output_dir, exist_ok=True)
    output_path = os.path.join(output_dir, f"{template_id}_manuscript.docx")
    doc.save(output_path)

    elapsed = (time.time() - start) * 1000
    logger.info(f"Formatted {template_id} in {elapsed:.0f}ms -> {output_path}")

    return output_path


def validate_manuscript(manuscript: StructuredManuscript, template_id: str) -> Dict[str, Any]:
    """Validate a manuscript against template requirements."""
    issues = []
    warnings = []
    score = 100.0

    if not manuscript.title:
        issues.append("Missing title")
        score -= 15

    if not manuscript.abstract:
        issues.append("Missing abstract")
        score -= 15

    if not manuscript.sections:
        issues.append("No sections found")
        score -= 20

    if not manuscript.references:
        warnings.append("No references found")
        score -= 10

    if not manuscript.authors:
        warnings.append("No authors listed")
        score -= 5

    config = TEMPLATE_CONFIGS.get(template_id)
    if config and config.requires_keywords and not manuscript.keywords:
        warnings.append(f"{config.name} requires keywords")
        score -= 5

    if manuscript.abstract:
        abstract_words = len(manuscript.abstract.split())
        if template_id in ("ieee", "springer") and abstract_words > 250:
            warnings.append(f"Abstract too long ({abstract_words} words, max 250)")
        elif template_id in ("acm", "elsevier") and abstract_words > 300:
            warnings.append(f"Abstract too long ({abstract_words} words, max 300)")

    score = max(0.0, min(100.0, score))

    return {
        "is_valid": len(issues) == 0,
        "issues": issues,
        "warnings": warnings,
        "score": round(score, 1),
    }
