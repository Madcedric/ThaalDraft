from typing import Any, Dict, List, Optional
import os
import time
from .schema import (
    FormatTemplate,
    FormatValidation,
    FormattedOutput,
    ExportType,
)
from .templates import get_template


def _validate_structured_data(
    structured_data: Dict[str, Any],
    template: FormatTemplate,
) -> FormatValidation:
    issues: List[str] = []
    warnings: List[str] = []

    title = structured_data.get("title", "")
    abstract = structured_data.get("abstract", "")
    sections = structured_data.get("sections", [])
    references = structured_data.get("references", [])
    authors = structured_data.get("authors", [])
    keywords = structured_data.get("keywords", [])

    if not title:
        issues.append("Missing title")
    elif template.title_font and template.title_font.size_pt:
        title_words = len(title.split())
        if title_words > 30:
            warnings.append(f"Title is very long ({title_words} words)")

    if not abstract:
        issues.append("Missing abstract")
    else:
        abstract_words = len(abstract.split())
        if template.id in ("ieee", "springer") and abstract_words > 250:
            warnings.append(f"Abstract too long for {template.name} ({abstract_words} words)")
        elif template.id in ("acm", "elsevier") and abstract_words > 300:
            warnings.append(f"Abstract too long for {template.name} ({abstract_words} words)")

    if not sections:
        issues.append("No sections found")
    if not references:
        warnings.append("No references found")
    if not authors:
        warnings.append("No authors listed")
    if template.requires_keywords and not keywords:
        warnings.append(f"{template.name} requires keywords")

    score = 100.0
    score -= len(issues) * 15
    score -= len(warnings) * 5
    score = max(0.0, min(100.0, score))

    return FormatValidation(
        is_valid=len(issues) == 0,
        issues=issues,
        warnings=warnings,
        score=round(score, 1),
    )


def _get_authors_list(structured_data: Dict) -> List[str]:
    authors = structured_data.get("authors", [])
    if isinstance(authors, list) and len(authors) > 0:
        if isinstance(authors[0], dict):
            return [a.get("name", str(a)) for a in authors]
        return [str(a) for a in authors]
    return []


def _build_ieee_docx(structured_data: Dict, template: FormatTemplate) -> Any:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(template.margins.top_inches)
        section.bottom_margin = Inches(template.margins.bottom_inches)
        section.left_margin = Inches(template.margins.left_inches)
        section.right_margin = Inches(template.margins.right_inches)

    style = doc.styles["Normal"]
    style.font.name = template.body_font.name
    style.font.size = Pt(template.body_font.size_pt)

    if structured_data.get("title"):
        p = doc.add_paragraph(structured_data["title"])
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0] if p.runs else p.add_run()
        run.font.name = template.title_font.name
        run.font.size = Pt(template.title_font.size_pt)
        run.bold = template.title_font.bold

    author_names = _get_authors_list(structured_data)
    if author_names:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(", ".join(author_names)).font.size = Pt(11)

    if structured_data.get("abstract"):
        doc.add_paragraph()
        p = doc.add_paragraph()
        r1 = p.add_run("Abstract\u2014")
        r1.bold = True
        r1.italic = True
        r2 = p.add_run(structured_data["abstract"])
        r2.bold = True

    if structured_data.get("keywords") and template.requires_keywords:
        p = doc.add_paragraph()
        r1 = p.add_run(f"{template.keywords_label}: ")
        r1.bold = True
        p.add_run(", ".join(structured_data["keywords"]))

    roman_numerals = ["I", "II", "III", "IV", "V", "VI", "VII", "VIII", "IX", "X"]
    for idx, section in enumerate(structured_data.get("sections", [])):
        heading = section.get("heading", "") if isinstance(section, dict) else ""
        if heading:
            roman = roman_numerals[idx] if idx < len(roman_numerals) else str(idx + 1)
            p = doc.add_paragraph(f"{roman}. {heading}")
            if p.runs:
                p.runs[0].bold = True
                p.runs[0].font.size = Pt(template.headings[0].font_size_pt)

        content = section.get("content", "") if isinstance(section, dict) else ""
        if content:
            p = doc.add_paragraph(content)
            p.paragraph_format.first_line_indent = Inches(0.14)

    if structured_data.get("references"):
        doc.add_paragraph(template.references_label.upper())
        refs = structured_data["references"]
        for idx, ref in enumerate(refs):
            ref_text = ref.get("raw_text", str(ref)) if isinstance(ref, dict) else str(ref)
            p = doc.add_paragraph(f"[{idx+1}] {ref_text}")
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.first_line_indent = Inches(-0.25)

    return doc


def _build_acm_docx(structured_data: Dict, template: FormatTemplate) -> Any:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(template.margins.top_inches)
        section.bottom_margin = Inches(template.margins.bottom_inches)
        section.left_margin = Inches(template.margins.left_inches)
        section.right_margin = Inches(template.margins.right_inches)

    style = doc.styles["Normal"]
    style.font.name = template.body_font.name
    style.font.size = Pt(template.body_font.size_pt)
    style.paragraph_format.space_after = Pt(6)

    if structured_data.get("title"):
        p = doc.add_paragraph(structured_data["title"])
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0] if p.runs else p.add_run()
        run.font.name = template.title_font.name
        run.font.size = Pt(template.title_font.size_pt)
        run.bold = template.title_font.bold

    author_names = _get_authors_list(structured_data)
    if author_names:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(", ".join(author_names)).font.size = Pt(10)

    if structured_data.get("abstract"):
        doc.add_paragraph()
        p = doc.add_paragraph()
        r1 = p.add_run("Abstract: ")
        r1.bold = True
        r2 = p.add_run(structured_data["abstract"])
        r2.italic = True
        r2.font.size = Pt(9)

    if structured_data.get("keywords") and template.requires_keywords:
        p = doc.add_paragraph()
        r1 = p.add_run(f"{template.keywords_label}: ")
        r1.bold = True
        p.add_run(", ".join(structured_data["keywords"]))

    for idx, section in enumerate(structured_data.get("sections", [])):
        heading = section.get("heading", "") if isinstance(section, dict) else ""
        if heading:
            p = doc.add_paragraph(heading)
            if p.runs:
                p.runs[0].bold = True
                p.runs[0].font.size = Pt(template.headings[0].font_size_pt)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)

        content = section.get("content", "") if isinstance(section, dict) else ""
        if content:
            doc.add_paragraph(content)

    if structured_data.get("references"):
        doc.add_paragraph(template.references_label)
        refs = structured_data["references"]
        for idx, ref in enumerate(refs):
            ref_text = ref.get("raw_text", str(ref)) if isinstance(ref, dict) else str(ref)
            p = doc.add_paragraph(f"{idx+1}. {ref_text}")
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.first_line_indent = Inches(-0.5)

    return doc


def _build_apa_docx(structured_data: Dict, template: FormatTemplate) -> Any:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(template.margins.top_inches)
        section.bottom_margin = Inches(template.margins.bottom_inches)
        section.left_margin = Inches(template.margins.left_inches)
        section.right_margin = Inches(template.margins.right_inches)

    style = doc.styles["Normal"]
    style.font.name = template.body_font.name
    style.font.size = Pt(template.body_font.size_pt)
    style.paragraph_format.line_spacing = 2.0

    if structured_data.get("title"):
        p = doc.add_paragraph(structured_data["title"])
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0] if p.runs else p.add_run()
        run.bold = True
        run.font.size = Pt(template.title_font.size_pt)

    author_names = _get_authors_list(structured_data)
    if author_names:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(", ".join(author_names)).font.size = Pt(12)

    if structured_data.get("abstract"):
        doc.add_paragraph()
        p = doc.add_paragraph()
        r1 = p.add_run("Abstract")
        r1.bold = True
        doc.add_paragraph(structured_data["abstract"])

    for idx, section in enumerate(structured_data.get("sections", [])):
        heading = section.get("heading", "") if isinstance(section, dict) else ""
        if heading:
            p = doc.add_paragraph(heading)
            if p.runs:
                p.runs[0].bold = True
                p.runs[0].font.size = Pt(template.headings[0].font_size_pt)
                p.runs[0].underline = True
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)

        content = section.get("content", "") if isinstance(section, dict) else ""
        if content:
            p = doc.add_paragraph(content)
            p.paragraph_format.first_line_indent = Inches(0.5)

    if structured_data.get("references"):
        doc.add_paragraph()
        p = doc.add_paragraph(template.references_label)
        if p.runs:
            p.runs[0].bold = True
            p.runs[0].font.size = Pt(14)

        refs = structured_data["references"]
        for ref in refs:
            ref_text = ref.get("raw_text", str(ref)) if isinstance(ref, dict) else str(ref)
            p = doc.add_paragraph(ref_text)
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.first_line_indent = Inches(-0.5)

    return doc


def _build_springer_docx(structured_data: Dict, template: FormatTemplate) -> Any:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(template.margins.top_inches)
        section.bottom_margin = Inches(template.margins.bottom_inches)
        section.left_margin = Inches(template.margins.left_inches)
        section.right_margin = Inches(template.margins.right_inches)

    style = doc.styles["Normal"]
    style.font.name = template.body_font.name
    style.font.size = Pt(template.body_font.size_pt)

    if structured_data.get("title"):
        p = doc.add_paragraph(structured_data["title"])
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0] if p.runs else p.add_run()
        run.font.name = template.title_font.name
        run.font.size = Pt(template.title_font.size_pt)
        run.bold = template.title_font.bold

    author_names = _get_authors_list(structured_data)
    if author_names:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(", ".join(author_names)).font.size = Pt(10)

    if structured_data.get("abstract"):
        doc.add_paragraph()
        p = doc.add_paragraph()
        r1 = p.add_run("Abstract ")
        r1.bold = True
        p.add_run(structured_data["abstract"])

    if structured_data.get("keywords") and template.requires_keywords:
        p = doc.add_paragraph()
        r1 = p.add_run(f"{template.keywords_label}: ")
        r1.bold = True
        p.add_run(", ".join(structured_data["keywords"]))

    for idx, section in enumerate(structured_data.get("sections", [])):
        heading = section.get("heading", "") if isinstance(section, dict) else ""
        if heading:
            p = doc.add_paragraph(heading)
            if p.runs:
                p.runs[0].bold = True
                p.runs[0].font.size = Pt(template.headings[0].font_size_pt)

        content = section.get("content", "") if isinstance(section, dict) else ""
        if content:
            doc.add_paragraph(content)

    if structured_data.get("references"):
        doc.add_paragraph(template.references_label)
        refs = structured_data["references"]
        for idx, ref in enumerate(refs):
            ref_text = ref.get("raw_text", str(ref)) if isinstance(ref, dict) else str(ref)
            p = doc.add_paragraph(f"{idx+1}. {ref_text}")
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.first_line_indent = Inches(-0.25)

    return doc


def _build_elsevier_docx(structured_data: Dict, template: FormatTemplate) -> Any:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(template.margins.top_inches)
        section.bottom_margin = Inches(template.margins.bottom_inches)
        section.left_margin = Inches(template.margins.left_inches)
        section.right_margin = Inches(template.margins.right_inches)

    style = doc.styles["Normal"]
    style.font.name = template.body_font.name
    style.font.size = Pt(template.body_font.size_pt)

    if structured_data.get("title"):
        p = doc.add_paragraph(structured_data["title"])
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0] if p.runs else p.add_run()
        run.font.name = template.title_font.name
        run.font.size = Pt(template.title_font.size_pt)
        run.bold = template.title_font.bold

    author_names = _get_authors_list(structured_data)
    if author_names:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(", ".join(author_names)).font.size = Pt(11)

    if structured_data.get("abstract"):
        doc.add_paragraph()
        p = doc.add_paragraph()
        r1 = p.add_run("Abstract\n")
        r1.bold = True
        p.add_run(structured_data["abstract"])

    if structured_data.get("keywords") and template.requires_keywords:
        p = doc.add_paragraph()
        r1 = p.add_run(f"{template.keywords_label}: ")
        r1.bold = True
        p.add_run("; ".join(structured_data["keywords"]))

    for idx, section in enumerate(structured_data.get("sections", [])):
        heading = section.get("heading", "") if isinstance(section, dict) else ""
        if heading:
            p = doc.add_paragraph(heading)
            if p.runs:
                p.runs[0].bold = True
                p.runs[0].font.size = Pt(template.headings[0].font_size_pt)

        content = section.get("content", "") if isinstance(section, dict) else ""
        if content:
            p = doc.add_paragraph(content)

    if structured_data.get("references"):
        doc.add_paragraph(template.references_label)
        refs = structured_data["references"]
        for idx, ref in enumerate(refs):
            ref_text = ref.get("raw_text", str(ref)) if isinstance(ref, dict) else str(ref)
            p = doc.add_paragraph(f"{idx+1}. {ref_text}")

    return doc


def _build_mla_docx(structured_data: Dict, template: FormatTemplate) -> Any:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(template.margins.top_inches)
        section.bottom_margin = Inches(template.margins.bottom_inches)
        section.left_margin = Inches(template.margins.left_inches)
        section.right_margin = Inches(template.margins.right_inches)

    style = doc.styles["Normal"]
    style.font.name = template.body_font.name
    style.font.size = Pt(template.body_font.size_pt)
    style.paragraph_format.line_spacing = 2.0

    if structured_data.get("title"):
        p = doc.add_paragraph(structured_data["title"])
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0] if p.runs else p.add_run()
        run.bold = True
        run.font.size = Pt(template.title_font.size_pt)

    author_names = _get_authors_list(structured_data)
    if author_names:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(", ".join(author_names)).font.size = Pt(12)

    if structured_data.get("abstract"):
        doc.add_paragraph()
        p = doc.add_paragraph()
        r1 = p.add_run("Abstract: ")
        r1.bold = True
        r2 = p.add_run(structured_data["abstract"])
        r2.italic = True

    for idx, section in enumerate(structured_data.get("sections", [])):
        heading = section.get("heading", "") if isinstance(section, dict) else ""
        if heading:
            p = doc.add_paragraph(heading)
            if p.runs:
                p.runs[0].bold = True
                p.runs[0].font.size = Pt(template.headings[0].font_size_pt)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)

        content = section.get("content", "") if isinstance(section, dict) else ""
        if content:
            p = doc.add_paragraph(content)
            p.paragraph_format.first_line_indent = Inches(0.5)

    if structured_data.get("references"):
        doc.add_paragraph()
        p = doc.add_paragraph(template.references_label)
        if p.runs:
            p.runs[0].bold = True
            p.runs[0].font.size = Pt(14)

        refs = structured_data["references"]
        for ref in refs:
            ref_text = ref.get("raw_text", str(ref)) if isinstance(ref, dict) else str(ref)
            p = doc.add_paragraph(ref_text)
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.first_line_indent = Inches(-0.5)

    return doc


DOCX_BUILDERS = {
    "ieee": _build_ieee_docx,
    "acm": _build_acm_docx,
    "springer": _build_springer_docx,
    "apa": _build_apa_docx,
    "mla": _build_mla_docx,
    "nature": _build_elsevier_docx,
}


def format_document(
    document_id: str,
    structured_data: Dict[str, Any],
    template_id: str,
    export_type: ExportType = ExportType.DOCX,
    output_dir: str = "exports",
) -> FormattedOutput:
    start_time = time.time()

    template = get_template(template_id)
    if template is None:
        raise ValueError(f"Template '{template_id}' not found")

    validation = _validate_structured_data(structured_data, template)

    if export_type == ExportType.DOCX:
        builder = DOCX_BUILDERS.get(template_id, _build_ieee_docx)
        doc = builder(structured_data, template)

        os.makedirs(output_dir, exist_ok=True)
        output_path = os.path.join(output_dir, f"{document_id}_{template_id}.docx")
        doc.save(output_path)

        processing_time_ms = (time.time() - start_time) * 1000

        return FormattedOutput(
            document_id=document_id,
            template_id=template_id,
            export_type=export_type,
            file_path=output_path,
            validation=validation,
            processing_metadata={
                "processing_time_ms": round(processing_time_ms, 2),
                "template_applied": template_id,
                "sections_formatted": len(structured_data.get("sections", [])),
            },
        )
    elif export_type == ExportType.PDF:
        temp_output = format_document(
            document_id=document_id,
            structured_data=structured_data,
            template_id=template_id,
            export_type=ExportType.DOCX,
            output_dir=output_dir,
        )

        try:
            from app.services.pdf_service import convert_docx_to_pdf
            pdf_path = os.path.join(output_dir, f"{document_id}_{template_id}.pdf")
            success = convert_docx_to_pdf(temp_output.file_path, pdf_path)
            if success:
                return FormattedOutput(
                    document_id=document_id,
                    template_id=template_id,
                    export_type=export_type,
                    file_path=pdf_path,
                    validation=validation,
                    processing_metadata=temp_output.processing_metadata,
                )
        except Exception:
            pass

        return temp_output
    else:
        raise ValueError(f"Export type '{export_type}' not supported")
