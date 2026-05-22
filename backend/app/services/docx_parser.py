import os
import uuid
import docx
import re
from fastapi import UploadFile, HTTPException

UPLOAD_DIR = "uploads"

# Ensure upload directory exists locally
os.makedirs(UPLOAD_DIR, exist_ok=True)

async def save_upload_file(upload_file: UploadFile) -> str:
    """Saves the uploaded file locally and returns the file path."""
    if not upload_file.filename.endswith('.docx'):
        raise HTTPException(status_code=400, detail="Invalid file type. Only .docx is supported.")
    
    file_id = str(uuid.uuid4())
    file_path = os.path.join(UPLOAD_DIR, f"{file_id}_{upload_file.filename}")
    
    with open(file_path, "wb") as buffer:
        content = await upload_file.read()
        buffer.write(content)
        
    return file_path

def parse_docx(file_path: str) -> dict:
    """Parses a docx file and extracts structured text using rule-based parsing."""
    doc = docx.Document(file_path)
    
    title = ""
    authors = []
    abstract = ""
    sections = []
    references = []
    tables_data = []
    figures_data = []
    
    current_section = "general"
    current_heading = ""
    
    # Regex for abstract, references, and figures
    abstract_pattern = re.compile(r'^(abstract|summary)$', re.IGNORECASE)
    reference_pattern = re.compile(r'^(references|bibliography|works cited)$', re.IGNORECASE)
    figure_pattern = re.compile(r'^(fig\.|figure)\s*\d+[:\.\-]', re.IGNORECASE)
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
            
        style_name = para.style.name.lower() if para.style else ""
        
        # Simple Title detection
        if not title and ("title" in style_name or len(text) < 150):
            title = text
            continue
            
        # Detect headings
        if "heading" in style_name or text.isupper():
            if abstract_pattern.match(text):
                current_section = "abstract"
                continue
            elif reference_pattern.match(text):
                current_section = "references"
                continue
            else:
                current_section = "body"
                current_heading = text
                sections.append({"heading": current_heading, "content": ""})
                continue
                
        # Detect figure captions
        if figure_pattern.match(text) or "caption" in style_name:
            figures_data.append(text)
            continue
            
        # Body logic
        if current_section == "abstract":
            abstract += text + "\n"
        elif current_section == "references":
            references.append(text)
        elif current_section == "body":
            if sections:
                sections[-1]["content"] += text + "\n"
            else:
                sections.append({"heading": "", "content": text + "\n"})
        else:
            if sections:
                sections[-1]["content"] += text + "\n"
            else:
                sections.append({"heading": "", "content": text + "\n"})
                current_section = "body"

    # Detect Tables
    for table in doc.tables:
        table_content = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            table_content.append(row_data)
        if table_content:
            tables_data.append(table_content)

    return {
        "title": title,
        "authors": authors,
        "abstract": abstract.strip(),
        "sections": sections,
        "references": references,
        "tables": tables_data,
        "figures": figures_data
    }
